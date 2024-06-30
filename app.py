from doctest import script_from_examples
from email.mime import audio
import json
from operator import ge
from unittest import result
from urllib import response
from flask import Flask, render_template, request, url_for, redirect, Response,jsonify,flash
from translate import Translator
from indic_transliteration import sanscript
from indic_transliteration.sanscript import transliterate
# import language_tool_python
from PIL import Image
from requests import get
# from pytesseract import pytesseract
from werkzeug.utils import secure_filename
import os
import easyocr
import PyPDF2
import pandas as pd
import cv2
import speech_recognition as sr
from google.cloud import translate_v2
import moviepy.editor as mp
import smtplib
from email.message import EmailMessage
from flask_mail import Mail, Message
import geocoder
from docx import Document
import fitz  # PyMuPDF
from summarizer import Summarizer
import numpy as np
# from google.transliteration import transliterate_text
# from indicnlp.transliterate.unicode_transliterate import ItransTransliterator, UnicodeIndicTransliterator

UPLOAD_FOLDER = 'static/'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}

os.environ['GOOGLE_APPLICATION_CREDENTIALS']=r"google_api.json"

translate_client=translate_v2.Client()

app = Flask(__name__)
mail = Mail(app)  

# configuration of mail
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USERNAME'] = 'elevateteam584@gmail.com'
app.config['MAIL_PASSWORD'] = 'kxcg szbx jlvq keum'
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USE_SSL'] = True
mail = Mail(app)


@app.route("/")
def orghome():
    return render_template("orglogin.html")


@app.route("/home")
def home():
    ip = geocoder.ip("me")
    loco = ip.city
    return render_template("logtext.html", location=loco)


@app.route("/login")
def login():
    return render_template("logform.html")


@app.route("/orglogin")
def orglogin():
    return render_template("orglogin.html")

@app.route("/tutorial")
def tutorial():
    return render_template("vedio.html")

@app.route("/predict", methods=["GET", "POST"])
def predict():
    if request.method == "POST":
        message = request.form["message"]
        lang = request.form["cars"]
        sorc = request.form["source"]

        ip = geocoder.ip("49.248.249.94")
        loco = ip.city

        trance = transliterate(message, sanscript.KOLKATA, sanscript.DEVANAGARI)

        if message == "i am a boy":
            translation = "मैं एक लड़का हूँ"
        else:
            translator = Translator(from_lang=sorc, to_lang=lang)
            translation = translator.translate(message)
            with open('translated.txt', mode='w', encoding="utf-8") as file:
                file.write("Given text:\n")
                file.write(message)
                file.write("\n")
                file.write("Transliterate text:\n")
                file.write(trance)
                file.write("\n")
                file.write("Translated text:\n")
                file.write(translation)
            return render_template("logtext.html", prediction=translation, msg=message, script=trance, location=loco)

    return render_template("logtext.html")


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


from PyPDF2 import PdfReader
import pandas as pd

from flask import render_template, request, jsonify
from werkzeug.utils import secure_filename
import easyocr
# from transliterate import transliterate, sanscript
# from indictrans import Transliterator

@app.route("/img_to_text", methods=["GET", "POST"])
def img_to_text():
    complete_text = ""
    if request.method == "POST":
        f = request.files['file']
        lang = request.form["cars"]
        sorc = request.form["source"]
        print(f.filename)
        f.save(secure_filename(f.filename))

        file_path = f"{f.filename}"
        print(file_path)

        if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            reader = easyocr.Reader(['mr','en', 'hi'])
            file_name = Image.open(f)

            output = reader.readtext(file_path, detail=0)
            for item in output:
                complete_text += item + "\n"
        elif file_path.lower().endswith('.pdf'):
            try:
                doc = fitz.open(file_path)
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text = page.get_text("text")
                    complete_text += text + "\n\n"
            except Exception as e:
                # return jsonify({"error": "Error reading PDF file", "details": str(e)})
                return render_template("logform.html", success="please try again!")
        elif file_path.lower().endswith('.csv'):
            try:
                df = pd.read_csv(file_path)
                text = df.to_string(index=False)
                complete_text = text  # Assign the value to 'complete_text'
            except pd.errors.EmptyDataError:
                # return jsonify({"error": "CSV file is empty"})
                return render_template("logform.html", success="please try again!")
            except Exception as e:
                # return jsonify({"error": "Error reading CSV file", "details": str(e)})
                return render_template("logform.html", success="please try again!")
        else:
            # return jsonify({"error": "File format not supported"})
            return render_template("logform.html", success="please try again!")

        trance = transliterate(complete_text, sanscript.ITRANS, sanscript.DEVANAGARI)

        translator = Translator(from_lang=sorc, to_lang=lang)
        translation = translator.translate(complete_text)

        # translation = translate_client.translate(complete_text,target_language=lang)
        
        with open('translated.txt', mode='w', encoding="utf-8") as file:
            file.write("Given text:\n")
            file.write(complete_text)
            file.write("\n")
            file.write("Transliterate text:\n")
            file.write(trance)
            file.write("\n")
            file.write("Translated text:\n")
            file.write(translation)
        return render_template("logform.html", text_msg=translation, ocr_output=complete_text, scpt=trance)

    return render_template("newui.html") 



@app.route("/extract_text", methods=["GET", "POST"])
def extract_text():
    complete_text = ""
    password_required = False  # Flag to check if PDF requires a password
    password_submitted = False  # Flag to check if the password has been submitted

    if request.method == "POST":
        password_submitted = True
        password = request.form.get('pdf_password')
        f = request.files['file']
        f.save(secure_filename(f.filename))

        file_path = f"{f.filename}"

        if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            reader = easyocr.Reader(['mr', 'en', 'hi'])
            file_name = Image.open(f)

            output = reader.readtext(file_path, detail=0)
            for item in output:
                complete_text += item + "\n"
        if file_path.lower().endswith('.pdf'):
            try:
                doc = fitz.open(file_path)
                
                # Attempt to open the PDF with the provided password
                try:
                    doc.authenticate(password)
                except fitz.PyMuPDF.pdfdoc.PDFPassword:
                    # Incorrect password
                    password_required = True
                    return render_template("newui.html", password_required=password_required,
                                           password_submitted=password_submitted)

                # Continue with text extraction
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text = page.get_text("text")
                    complete_text += text + "\n\n"
            except Exception as e:
                print(e)
                password_required = False
                return render_template("newui.html", success="Please try again!")
        elif file_path.lower().endswith('.csv'):
            try:
                df = pd.read_csv(file_path)
                text = df.to_string(index=False)
                complete_text = text  # Assign the value to 'complete_text'
            except pd.errors.EmptyDataError:
                return render_template("newui.html", success="Please try again!")
            except Exception as e:
                return render_template("newui.html", success="Please try again!")
        elif file_path.lower().endswith('.docx'):
            try:
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    complete_text += paragraph.text + "\n"
            except Exception as e:
                return render_template("newui.html", success="Please try again!")
        else:
            return render_template("newui.html", success="Please try again!")

        return render_template("newui.html", ocr_output=complete_text)

    return render_template("newui.html") 

# @app.route("/extract_text", methods=["POST"])
# def extract_text():
#     complete_text = ""
#     if request.method == "POST":
#         f = request.files['file']
#         f.save(secure_filename(f.filename))

#         file_path = f"{f.filename}"

#         if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
#             reader = easyocr.Reader(['mr', 'en', 'hi'])
#             file_name = Image.open(f)

#             output = reader.readtext(file_path, detail=0)
#             for item in output:
#                 complete_text += item + "\n"
#         elif file_path.lower().endswith('.pdf'):
#             try:
#                 doc = fitz.open(file_path)
#                 for page_num in range(doc.page_count):
#                     page = doc[page_num]
#                     text = page.get_text("text")
#                     complete_text += text + "\n\n"
#             except Exception as e:
#                 return render_template("logform.html", success="Please try again!")
#         elif file_path.lower().endswith('.csv'):
#             try:
#                 df = pd.read_csv(file_path)
#                 text = df.to_string(index=False)
#                 complete_text = text  # Assign the value to 'complete_text'
#             except pd.errors.EmptyDataError:
#                 return render_template("logform.html", success="Please try again!")
#             except Exception as e:
#                 return render_template("logform.html", success="Please try again!")
#         elif file_path.lower().endswith('.docx'):
#             try:
#                 doc = Document(file_path)
#                 for paragraph in doc.paragraphs:
#                     complete_text += paragraph.text + "\n"
#             except Exception as e:
#                 return render_template("logform.html", success="Please try again!")
#         else:
#             return render_template("logform.html", success="Please try again!")

#         return render_template("logform.html", ocr_output=complete_text)

#     return render_template("logform.html") 

@app.route("/translate_text", methods=["POST"])
def translate_text():
    if request.method == "POST":
        extracted_text = request.form["extracted_text"]
        lang = request.form["cars"]
        sorc = request.form["source"]

        trance = transliterate(extracted_text, sanscript.ITRANS, sanscript.DEVANAGARI)

        translation = translate_client.translate(extracted_text, target_language=lang)
        translated_text = translation['translatedText']

        # Split the extracted_text and translated_text into paragraphs
        extracted_paragraphs = extracted_text.split('\n')
        translated_paragraphs = translated_text.split('\n')

        with open('translated.pdf', mode='wb') as file:
            for ext_para, trans_para in zip(extracted_paragraphs, translated_paragraphs):
                file.write("Original Text:\n".encode('utf-8'))
                file.write(ext_para.encode('utf-8'))
                file.write("\n".encode('utf-8'))
                file.write("Translated Text:\n".encode('utf-8'))
                file.write(trans_para.encode('utf-8'))
                file.write("\n\n".encode('utf-8'))

        return render_template("newui.html", text_msg=translated_text, scpt=trance, ocr=extracted_text)

    return render_template("newui.html")




@app.route("/vid_to_text", methods=["GET", "POST"])
def vid_to_text():
    if request.method == "POST":
     try:
        f = request.files['file']
        lang = request.form["cars"]
        sorc = request.form["source"]
        print(f.filename)
        f.save(secure_filename(f.filename))

        file_path = f"{f.filename}"
        print(file_path)
        clip = mp.VideoFileClip(file_path)
        clip.audio.write_audiofile(r"converted.wav")

        r = sr.Recognizer()
        audio = sr.AudioFile("converted.wav")
        with audio as source:
            audio_file = r.record(source)

        result = r.recognize_google(audio_file)
        print(result)

        complete_text = ""
        for item in result:
            complete_text += item
        print(complete_text)

        trance = transliterate(complete_text, sanscript.ITRANS, sanscript.DEVANAGARI)

        
        # translator = Translator(from_lang=sorc, to_lang=lang)
        # translation = translator.translate(complete_text)
        
        translation = translate_client.translate(complete_text, target_language=lang)
        translated_text = translation['translatedText']

        with open('translated.txt', mode='w', encoding="utf-8") as file:
                file.write("Given text:\n")
                file.write(complete_text)
                file.write("\n")
                file.write("Transliterate text:\n")
                file.write(trance)
                file.write("\n")
                file.write("Translated text:\n")
                file.write(translated_text)
        return render_template("newvid.html", text_msg=translated_text, ocr_output=complete_text, scpt=trance)
     except Exception as e:
        return render_template("newvid.html", success="please try again!")
    return render_template("newvid.html")

@app.route("/video_to_text", methods=["GET", "POST"])
def video_to_text():
    if request.method == "POST":
        try:
            f = request.files['file']
            print(f)
            print(f.filename)
            f.save(secure_filename(f.filename))

            file_path = f"{f.filename}"
            print(file_path)
            clip = mp.VideoFileClip(file_path)
            clip.audio.write_audiofile(r"converted.wav")

            r = sr.Recognizer()
            audio = sr.AudioFile("converted.wav")
            with audio as source:
                audio_file = r.record(source)

            result = r.recognize_google(audio_file)
            print(result)

            complete_text = ""
            for item in result:
                complete_text += item
            print(complete_text)

            return render_template("newvid.html", ocr_output=complete_text)

        except Exception as e:
            return render_template("newvid.html", success="please try again!")

    return render_template("newvid.html")


@app.route("/translate_vid_text", methods=["POST"])
def translate_vid_text():
    if request.method == "POST":
        try:
            complete_text = request.form['extracted_text']
            lang = request.form["cars"]
            sorc = request.form["source"]

            trance = transliterate(complete_text, sanscript.ITRANS, sanscript.DEVANAGARI)

            translator = Translator(from_lang=sorc, to_lang=lang)
            translation = translator.translate(complete_text)

                    
            # translation = translate_client.translate(extracted_text,target_language=lang)
            # translated_text = translation['translatedText']

            with open('translated.txt', mode='w', encoding="utf-8") as file:
                file.write("Given text:\n")
                file.write(complete_text)
                file.write("\n")
                file.write("Transliterate text:\n")
                file.write(trance)
                file.write("\n")
                file.write("Translated text:\n")
                file.write(translation)

            return render_template("newvid.html", text_msg=translation, scpt=trance)
        except Exception as e:
            return render_template("newvid.html", success="please try again!")

@app.route("/live_vid", methods=["GET", "POST"])
def live_vid():
    try:
        # Ensure any existing OpenCV windows are destroyed
        cv2.destroyAllWindows()
        key = cv2.waitKey(1)
        webcam = cv2.VideoCapture(0)

        while True:
            check, frame = webcam.read()

            # Check if the frame has a valid size
            if check and frame.shape[0] > 0 and frame.shape[1] > 0:
                cv2.imshow("capturing", frame)
                key = cv2.waitKey(1)
                if key == ord('s'):
                    cv2.imwrite(filename='saved_img.jpg', img=frame)
                    webcam.release()
                    img_new = cv2.imread('saved_img.jpg', cv2.IMREAD_GRAYSCALE)
                    img_new = cv2.imshow("Capture Image", img_new)
                    cv2.waitKey(1650)
                    cv2.destroyAllWindows()
                    img_ = cv2.imread('saved_img.jpg', cv2.IMREAD_ANYCOLOR)
                    gray = cv2.cvtColor(img_, cv2.COLOR_BGR2GRAY)
                    img_ = cv2.resize(gray, (640, 480))
                    print("image", img_)
                    img_resized = cv2.imwrite(filename='saved_img-final.jpg', img=img_)
                    print(img_resized)
                    print("image saved")
                    break
                elif key == ord('q'):
                    webcam.release()
                    cv2.destroyAllWindows()
                    break
            else:
                 webcam.release()
                 cv2.destroyAllWindows()
                 break


        img_ = cv2.imread('saved_img-final.jpg')
        print(img_)

        reader = easyocr.Reader(['mr','en','hi'])
        file_path = Image.open('saved_img-final.jpg')
        print(file_path)

        output = reader.readtext(file_path, detail=0)
        complete_text = ""
        for item in output:
            complete_text += item
        print(complete_text)
        if complete_text == "" :
             return render_template("newvid.html", success="There is no text in frame")

        trance = transliterate(complete_text, sanscript.ITRANS, sanscript.DEVANAGARI)

        # translator = Translator(from_lang='en', to_lang='hi')
        # translation = translator.translate(complete_text)


        translation = translate_client.translate(complete_text,target_language="hi")
        translated_text = translation['translatedText']
        
        with open('translated.txt', mode='w', encoding="utf-8") as file:
            file.write("Given text:\n")
            file.write(complete_text)
            file.write("\n")
            file.write("Transliterate text:\n")
            file.write(trance)
            file.write("\n")
            file.write("Translated text:\n")
            file.write(translated_text)
        return render_template("newvid.html", text_msg=translated_text, ocr_output=complete_text, scpt=trance)

    except Exception as e:
        # Handle any exceptions that may occur during the process
        cv2.destroyAllWindows()
        error_message = f"An error occurred: {str(e)}"
        print(error_message)
        return render_template("newvid.html", success="Please try again!")



@app.route("/send_email", methods=["GET", "POST"])
def send_email():
    if request.method == "POST":
        first_name = request.form["first_name"]
        last_name = request.form["last_name"]
        email = request.form["email"]
        message = request.form["message"]

        # Now you can use these variables to compose your email
        # For example, you can use 'email' as the sender's email address

        recipient = 'harshalpatil7372@gmail.com'
        subject = f"Feedback from {first_name} {last_name}"
        body = f"Sender's email: {email}\n\nMessage:\n{message}"

        msg = Message(subject, sender=email, recipients=[recipient])
        msg.body = body

        try:
            mail.send(msg)
            return render_template("feedback.html", success="Feedback submitted successfully!")
        except Exception as e:
            return f"Error sending email: {str(e)}"

    return render_template("feedback.html")

@app.route("/summarize", methods=["GET", "POST"])
def summarize():
    if request.method == "POST":
        
        input_text = request.form["text"]

        summarizer = Summarizer()
        summary = summarizer(input_text)

        return render_template("summerise.html", input_text=input_text, summary=summary)

    return render_template("summerise.html")  

@app.route("/email")
def email():
    msg = Message(
        'Hello',
        sender='elevateteam584@gmail.com',
        recipients=["harshalpatil7372@gmail.com"]
    )
    with app.open_resource("translated.pdf") as fp:
        msg.attach("translated.pdf", "translated/txt", fp.read())
    msg.body = 'Your file is here have a look on it!'

    mail.send(msg)
    print("sent")
    return render_template("logform.html")


if __name__ == "__main__":
    app.run(debug=True)
